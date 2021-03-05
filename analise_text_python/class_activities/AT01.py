# Atividade em aula - AT01
# Autor: Newmar Wegner, Paulo Gamero, Kleberson Nascimento
# Date: 05/03/2021

import os
import docx
import nltk
from nltk.corpus import stopwords

# Open document with docx
arq = os.getcwd() + '/documento.docx'
doc = docx.Document(arq)

# Create a paragraphs list
paragraph_list = []
for p in doc.paragraphs:
    paragraph_list.append(p.text)

# Applying word tokenize
words = nltk.word_tokenize(' '.join(paragraph_list))

# Applying stopwords
stopword = stopwords.words('portuguese')
my_stop_w = [',', '[', ':', '\'', '.', '...',
             '?', ']', '!', '/', '-', "''"]
list_stopwords = [p for p in words if ((p not in stopword) and (p not in my_stop_w))]

# 10 words most frequentily
pos_tag = nltk.pos_tag(list_stopwords)
fd_ngramas = nltk.FreqDist(pos_tag)

# create bigrams
bigrama = nltk.FreqDist([ng for ng in nltk.ngrams(list_stopwords, 2)])

# select first paragraph
first_paragraph = []
for p in doc.paragraphs:
        first_paragraph.append(p.text)

first_paragraph = first_paragraph[0]


# generating pos
stopwords = nltk.corpus.stopwords.words('portuguese')
words_fp = nltk.word_tokenize(first_paragraph)
list_word_fp = [p for p in words_fp if ((p not in stopword) and (p not in my_stop_w))]

# generating ner
entities = []
labels = []
chuncks = nltk.ne_chunk(nltk.pos_tag(list_word_fp), binary=True)

for chunck in chuncks:
    if hasattr(chunck, 'label'):
        entities.append(' '.join(c[0] for c in chunck))
        labels.append(chunck.label())
    
if __name__ == '__main__':
    print(f'Questão 01: O número de palavras no texto é {len(list_stopwords)}.')
    print(f'Questão 02: As 10 palavras mais utilizadas no texto são: {[v[0][0] for v in fd_ngramas.most_common(10)]}')
    print(f'Questão 03: Os 10 bigramas mais comuns são: {[v[0] for v in bigrama.most_common(10)]}')
    print(f"Questão 04: O texto possui {len(nltk.sent_tokenize((' '.join(words))))} sentenças.")
    print(f'Questão 05: As palavras e suas classificações POS são: {nltk.pos_tag(list_word_fp)}')
    print(f'Questão 06: As palavras e suas classificações NER são: {list(set(zip(entities, labels)))}')
